from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Page margins ---
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(1)
section.top_margin  = section.bottom_margin = Inches(1)

NAVY  = RGBColor(0x0f, 0x28, 0x50)
BLUE  = RGBColor(0x1d, 0x4e, 0xd8)
GRAY  = RGBColor(0x6b, 0x72, 0x80)
WHITE = RGBColor(0xff, 0xff, 0xff)
GOLD  = RGBColor(0xc9, 0xa8, 0x4c)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18 if level == 1 else 14)
    run.font.color.rgb = color or (NAVY if level == 1 else BLUE)
    return p

def add_body(doc, text, bold_prefix=None, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    if bold_prefix:
        run_b = p.add_run(bold_prefix + "  ")
        run_b.bold = True
        run_b.font.size = Pt(11)
        run_b.font.color.rgb = NAVY
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_step(doc, number, title, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0)
    r1 = p.add_run(f"Step {number}: ")
    r1.bold = True; r1.font.size = Pt(12); r1.font.color.rgb = BLUE
    r2 = p.add_run(title)
    r2.bold = True; r2.font.size = Pt(12); r2.font.color.rgb = NAVY
    b = doc.add_paragraph()
    b.paragraph_format.left_indent = Inches(0.3)
    b.paragraph_format.space_after = Pt(4)
    run = b.add_run(body)
    run.font.size = Pt(11)

def add_box(doc, emoji, title, body_lines):
    """Shaded tip box using a 1-col table"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    # shading
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EFF6FF')
    tcPr.append(shd)
    # border
    for side in ('top','bottom','left','right'):
        bd = OxmlElement(f'w:top' if side=='top' else f'w:{side}')
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    p = cell.add_paragraph()
    r = p.add_run(f"{emoji}  {title}")
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = BLUE
    for line in body_lines:
        bp = cell.add_paragraph()
        bp.paragraph_format.left_indent = Inches(0.15)
        br = bp.add_run(line)
        br.font.size = Pt(11)
    doc.add_paragraph()  # spacing after

def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1d4ed8')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(12)

# ══════════════════════════════════════════
#  COVER
# ══════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run("🏢  LaVau & Lucas Pipeline")
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Your Simple CRM User Guide")
r2.font.size = Pt(16); r2.font.color.rgb = BLUE; r2.bold = True

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Hawaii · California · Arizona  |  Commercial · Land · Business")
r3.font.size = Pt(11); r3.font.color.rgb = GRAY

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run("Written for Frank Lucas  —  Easy enough for a 5th grader 🤙")
r4.font.size = Pt(11); r4.font.color.rgb = GRAY; r4.italic = True

doc.add_page_break()

# ══════════════════════════════════════════
#  SECTION 1: WHAT IS THIS?
# ══════════════════════════════════════════
add_heading(doc, "1.  What Is This Thing?")
add_divider(doc)
add_body(doc, "Think of the LaVau & Lucas Pipeline like a digital sticky-note board for every deal you and Sean are working on. Instead of hunting through Apple Notes, texts, and emails — everything lives in one place, organized by what's happening with each deal right now.")
doc.add_paragraph()
add_body(doc, "It's a single file on your computer (LaVau-Lucas-CRM.html). You open it in Safari or Chrome like a website — no apps to install, no accounts to create, no internet needed.")
doc.add_paragraph()
add_body(doc, "🟦 BLUE badge = Commercial deal", indent=True)
add_body(doc, "🟢 GREEN badge = Land deal", indent=True)
add_body(doc, "🟣 PURPLE badge = Business deal", indent=True)

doc.add_paragraph()

# ══════════════════════════════════════════
#  SECTION 2: TWO VIEWS
# ══════════════════════════════════════════
add_heading(doc, "2.  Two Views — Yours & Sean's")
add_divider(doc)
add_body(doc, "The pipeline has two modes. Think of it like a light switch:")
doc.add_paragraph()
add_body(doc, "Sean's View (Default):", bold_prefix="👁")
add_body(doc, "This is what Sean sees when he opens the file. Clean, read-only. He can see all the deals, statuses, contacts, and next actions — but he can't accidentally change anything.", indent=True)
doc.add_paragraph()
add_body(doc, "Your View — Manage Mode:", bold_prefix="⚙️")
add_body(doc, "Tap the '⚙ Manage' button in the top-right corner. Now you can add new deals, edit existing ones, and update anything. Tap it again to exit.", indent=True)

doc.add_paragraph()

# ══════════════════════════════════════════
#  SECTION 3: ADDING A NEW DEAL
# ══════════════════════════════════════════
add_heading(doc, "3.  How to Add a New Deal")
add_divider(doc)
add_step(doc, 1, "Open the file", "Double-click LaVau-Lucas-CRM.html on your computer. It opens in your web browser.")
add_step(doc, 2, "Enter Manage Mode", "Tap the '⚙ Manage' button in the top-right corner of the screen. It turns white when active.")
add_step(doc, 3, "Tap '+ Add Deal'", "A blue button appears in the bottom-right corner. Tap it to open a form.")
add_step(doc, 4, "Fill in the form", "Fill in what you know — you don't need to fill in everything. At minimum: Deal Name, Type, and Status.")
add_step(doc, 5, "Tap 'Save Deal'", "The deal appears instantly on the board.")
add_step(doc, 6, "Save & Share", "Tap '💾 Save & Download' at the bottom. This downloads a fresh copy of the file with your changes baked in. Replace the old file with this new one.")

doc.add_paragraph()

# ══════════════════════════════════════════
#  SECTION 4: EDITING A DEAL
# ══════════════════════════════════════════
add_heading(doc, "4.  How to Edit or Update a Deal")
add_divider(doc)
add_step(doc, 1, "Enter Manage Mode", "Tap '⚙ Manage' in the top-right.")
add_step(doc, 2, "Find the deal", "Scroll to the deal card you want to update.")
add_step(doc, 3, "Tap 'Edit'", "A blue Edit button appears on the bottom-right of each card in Manage Mode.")
add_step(doc, 4, "Update the fields", "Change whatever needs updating — status, last contact date, next action, notes, anything.")
add_step(doc, 5, "Tap 'Save'", "Changes are saved instantly.")
add_step(doc, 6, "Download & Replace", "Tap '💾 Save & Download' to lock in your changes for Sean.")

doc.add_paragraph()

# ══════════════════════════════════════════
#  SECTION 5: THE FIELDS EXPLAINED
# ══════════════════════════════════════════
add_heading(doc, "5.  What Each Field Means")
add_divider(doc)

fields = [
    ("Deal / Project Name", "A short nickname for the deal. Example: 'Barrett Winery' or 'ROE Visual Expansion'"),
    ("Deal Type", "Commercial, Land, or Business. Picks the color of the badge on the card."),
    ("Status", "Where the deal is in your pipeline right now. Choose the one that fits best (see Section 6 below)."),
    ("Property Address", "The street address of the property. Leave blank if still searching."),
    ("Asking Price / Offer Price", "The list price and what you offered or are under contract at."),
    ("Client / Seller Name", "The person you're working with on this deal."),
    ("Client Phone", "Their phone number. On iPhone, tapping this calls or texts them directly."),
    ("Client Email", "Their email. Tapping it opens a new email to them."),
    ("Last Contacted", "The date you last talked to this person. Keeps track of how long since you reached out."),
    ("Property Details", "The building specs — square footage, zoning, ceiling height, docks, power, etc. Separate each spec with a | symbol and they become individual tags on the card."),
    ("Property Profile Link", "A link to the LoopNet, CoStar, or Crexi listing. Tap 'View Listing' on the card to open it."),
    ("Next Action", "The ONE thing that needs to happen next to move this deal forward."),
    ("Next Action Due Date", "When that action needs to happen by. If it passes, the card turns red and floats to the top."),
    ("Notes", "Everything else — context, deal structure, things to remember, personality notes about the client."),
    ("Archive", "Moves a deal off the main board without deleting it. Use for deals that are stalled but not dead."),
]

for name, desc in fields:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"{name}:  ")
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = NAVY
    r2 = p.add_run(desc)
    r2.font.size = Pt(11)

doc.add_paragraph()

# ══════════════════════════════════════════
#  SECTION 6: STATUSES
# ══════════════════════════════════════════
add_heading(doc, "6.  Deal Statuses — The Pipeline")
add_divider(doc)
add_body(doc, "Each deal moves through stages from left to right. Pick the one that best describes where you are today:")
doc.add_paragraph()

statuses = [
    ("🔵 Prospect",       "You know about this opportunity but haven't officially started working it yet."),
    ("🟡 Active Pursuit", "You're actively working this deal — making calls, doing research, building the relationship."),
    ("🟠 Offer / LOI",    "You've submitted a Letter of Intent or formal offer. Waiting on a response."),
    ("🔴 Under Contract", "Both sides have signed. The deal is locked in, now you execute."),
    ("🟣 Due Diligence",  "Inspections, title work, financing, zoning checks — all the homework before closing."),
    ("🟢 Closing",        "Almost done. Final paperwork, funding, and transfer happening."),
    ("✅ Closed",         "Done\! The deal is complete. It moves to the 'Closed' section at the bottom."),
    ("⚫ Inactive",       "The deal fell apart or is on hold. It moves off the main board."),
]

for status, desc in statuses:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.2)
    r1 = p.add_run(f"{status}  —  ")
    r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.size = Pt(11)

doc.add_paragraph()

# ══════════════════════════════════════════
#  SECTION 7: SYNCING WITH SEAN
# ══════════════════════════════════════════
add_heading(doc, "7.  Sharing Updates with Sean")
add_divider(doc)
add_body(doc, "This is the most important habit to build. Here's the routine:")
doc.add_paragraph()
add_step(doc, 1, "You make changes", "Add a deal, update a status, log a contact date — whatever happened today.")
add_step(doc, 2, "Tap '💾 Save & Download'", "This creates a fresh version of the file with all your changes locked in. It downloads automatically.")
add_step(doc, 3, "Replace the old file", "Go to iCloud Drive (or wherever you both share the file). Delete the old LaVau-Lucas-CRM.html and drag in the new one.")
add_step(doc, 4, "Sean opens it", "Next time Sean taps the file, he sees everything updated. No texts needed, no calls, no explaining.")

doc.add_paragraph()
add_body(doc, "💡 Tip: If you host the file on Netlify (free), Sean gets a real link like https://lavau-lucas.netlify.app — no file sharing needed. You just re-upload the new file to Netlify each time.")

doc.add_paragraph()

# ══════════════════════════════════════════
#  SECTION 8: HOSTING ON NETLIFY
# ══════════════════════════════════════════
add_heading(doc, "8.  Going Live on the Internet (Netlify) — 2 Minutes")
add_divider(doc)
add_body(doc, "Netlify gives your pipeline a real web address so Sean (and anyone you want) can open it from their phone without needing the file. It's free and takes 2 minutes.")
doc.add_paragraph()
add_step(doc, 1, "Go to the website", "Open netlify.com/drop in your browser on your Mac.")
add_step(doc, 2, "Drag your file", "Find LaVau-Lucas-CRM.html on your computer and drag it onto the Netlify page. Drop it in the big box that says 'Drag and drop your site folder here.'")
add_step(doc, 3, "Get your link", "Netlify instantly gives you a link like: https://random-name-123.netlify.app. That's your live pipeline.")
add_step(doc, 4, "Share with Sean", "Text Sean that link. He opens it in Safari on his iPhone. Done.")
add_step(doc, 5, "Update it later", "When you download a new version of the file, go back to Netlify, sign in (free), find your site, and drag the new file in. Sean's link stays the same — it just shows the updated content.")

doc.add_paragraph()

# ══════════════════════════════════════════
#  SECTION 9: CONNECTING TO GITHUB
# ══════════════════════════════════════════
add_heading(doc, "9.  Connecting to GitHub (Optional — For Later)")
add_divider(doc)
add_body(doc, "GitHub is like a file locker that also tracks every change you ever make. It's more advanced than Netlify Drop, but worth learning when you're ready. Here's the simple version:")
doc.add_paragraph()
add_step(doc, 1, "Create a free account", "Go to github.com and sign up. It's free.")
add_step(doc, 2, "Create a repository", "Think of this like a project folder online. Click the green 'New' button. Name it 'lavau-lucas-pipeline'. Check the box to 'Add a README file'. Click 'Create repository'.")
add_step(doc, 3, "Upload your file", "Click 'Add file' → 'Upload files'. Drag your LaVau-Lucas-CRM.html in. Click 'Commit changes'.")
add_step(doc, 4, "Turn on GitHub Pages", "Go to Settings → Pages. Under 'Source', pick 'main' branch and click Save. GitHub gives you a free public link like: https://frankalemanylucas.github.io/lavau-lucas-pipeline")
add_step(doc, 5, "Share that link", "Same as Netlify — text it to Sean. Every time you re-upload the file, the link automatically shows the new version.")
doc.add_paragraph()
add_body(doc, "💡 Netlify Drop is faster for tomorrow. GitHub Pages is better long-term because it keeps a history of every version you've ever saved.", indent=True)

doc.add_paragraph()

# ══════════════════════════════════════════
#  SECTION 10: GMAIL + MESSAGES
# ══════════════════════════════════════════
add_heading(doc, "10.  Connecting Gmail & Your Phone")
add_divider(doc)
add_body(doc, "You can connect Gmail directly to this tool (Claude / Cowork) so that I can help you draft follow-up emails to clients, search your inbox for past conversations, and log activity. Here's how:")
doc.add_paragraph()
add_body(doc, "Gmail:", bold_prefix="📧")
add_body(doc, "In Cowork, look for the 'Connect' button in the bottom toolbar or settings. Search for Gmail and connect your Google account. Once connected, you can say things like: 'Draft a follow-up email to John Barrett about his paperwork' — and I'll write it for you.", indent=True)
doc.add_paragraph()
add_body(doc, "iMessage:", bold_prefix="💬")
add_body(doc, "iMessage doesn't have a direct connector yet (Apple keeps it locked tight). The best workaround: when you're in a conversation with a client in Messages, copy the key info and paste it here — I'll help you log it, update the deal, and write a response.", indent=True)
doc.add_paragraph()
add_body(doc, "Your PC:", bold_prefix="💻")
add_body(doc, "If your pipeline is hosted on Netlify or GitHub Pages, you can open it on any PC just by going to the link in any browser. No file transfers needed. If you want to edit deals from your PC, you'll need the HTML file on that computer too — easiest way is to email it to yourself or put it in Google Drive.", indent=True)

doc.add_paragraph()

# ══════════════════════════════════════════
#  SECTION 11: QUICK REFERENCE CHEAT SHEET
# ══════════════════════════════════════════
add_heading(doc, "11.  Quick Reference Cheat Sheet")
add_divider(doc)

cheat = [
    ("Open the pipeline",        "Double-click LaVau-Lucas-CRM.html  OR  go to your Netlify/GitHub link"),
    ("Enter edit mode",          "Tap '⚙ Manage' in the top-right corner"),
    ("Add a new deal",           "Enter Manage mode → tap blue '+ Add Deal' button (bottom-right)"),
    ("Edit a deal",              "Enter Manage mode → tap 'Edit' on the deal card"),
    ("Mark a deal inactive",     "Enter Manage mode → tap 'Archive' on the deal card"),
    ("Update status",            "Edit the deal → change the Status dropdown"),
    ("Log last contact",         "Edit the deal → fill in 'Last Contacted' with today's date"),
    ("Add property specs",       "Edit the deal → type specs in 'Property Details' separated by |"),
    ("Link to a listing",        "Edit the deal → paste LoopNet/CoStar URL in 'Property Profile Link'"),
    ("Save & share with Sean",   "Tap '💾 Save & Download' → replace the old file in iCloud / Netlify"),
    ("See only overdue deals",   "Tap 'Overdue' in the filter bar at the top"),
    ("See only Business deals",  "Tap 'Business' in the filter bar"),
    ("View closed deals",        "Scroll down → tap 'Show Closed / Inactive'"),
]

tbl = doc.add_table(rows=len(cheat)+1, cols=2)
tbl.style = 'Table Grid'

# Header row
hdr = tbl.rows[0]
for i, txt in enumerate(["What You Want to Do", "How to Do It"]):
    cell = hdr.cells[i]
    cell.width = Inches(2.5 if i==0 else 5.0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '0F2850')
    tcPr.append(shd)
    p = cell.paragraphs[0]
    run = p.add_run(txt)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = WHITE

for r_idx, (action, how) in enumerate(cheat):
    row = tbl.rows[r_idx+1]
    fill = 'EFF6FF' if r_idx % 2 == 0 else 'FFFFFF'
    for c_idx, txt in enumerate([action, how]):
        cell = row.cells[c_idx]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(10)
        if c_idx == 0:
            run.bold = True
            run.font.color.rgb = NAVY

doc.add_paragraph()

# ══════════════════════════════════════════
#  FOOTER / CLOSING
# ══════════════════════════════════════════
p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_end.paragraph_format.space_before = Pt(20)
r_end = p_end.add_run("LaVau & Lucas Pipeline  ·  Built by Claude for Frank & Sean  🤙")
r_end.font.size = Pt(10)
r_end.font.color.rgb = GRAY
r_end.italic = True

doc.save('/sessions/vibrant-nice-brahmagupta/mnt/outputs/LaVau-Lucas-User-Guide.docx')
print("✅ Done")
